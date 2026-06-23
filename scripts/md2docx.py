# -*- coding: utf-8 -*-
"""把 方案说明.md 转为排版精致的 DOCX"""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "方案说明.md"
OUT = "方案说明.docx"

BRAND = RGBColor(0x1D, 0x4E, 0xD8)   # 主题蓝
DARK = RGBColor(0x1F, 0x29, 0x37)    # 正文深灰
GRAY = RGBColor(0x6B, 0x72, 0x80)    # 副标题灰

doc = Document()

# 页边距
sec = doc.sections[0]
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.4)
sec.left_margin = Cm(2.8)
sec.right_margin = Cm(2.8)

# 正文默认样式：宋体 + 西文 Times New Roman
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(11)
normal.font.color.rgb = DARK
normal.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def set_cn_font(run, font="宋体"):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def add_runs_with_bold(p, text):
    """解析 **加粗**，加粗部分用主题色"""
    text = text.replace("`", "")
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2])
            r.bold = True
            r.font.color.rgb = BRAND
            set_cn_font(r)
        else:
            r = p.add_run(part)
            set_cn_font(r)


def add_bottom_border(p, color="1D4ED8", size="6"):
    """给段落加底部边框（用作标题下划线/分隔线）"""
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def add_cover(title, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = BRAND
    set_cn_font(r, "黑体")

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp.paragraph_format.space_after = Pt(16)
    sr = sp.add_run(subtitle)
    sr.font.size = Pt(11)
    sr.font.color.rgb = GRAY
    set_cn_font(sr, "楷体")
    add_bottom_border(sp, "C7D2FE", "4")


def add_h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = BRAND
    set_cn_font(r, "黑体")
    add_bottom_border(p)


def add_body(text):
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.space_after = Pt(6)
    fmt.first_line_indent = Pt(22)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_runs_with_bold(p, text)


def add_list_item(text, ordered, idx=0):
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.space_after = Pt(4)
    fmt.left_indent = Pt(20)
    fmt.first_line_indent = Pt(-2)
    prefix = f"{idx}. " if ordered else "• "
    rp = p.add_run(prefix)
    rp.bold = ordered
    if ordered:
        rp.font.color.rgb = BRAND
    set_cn_font(rp)
    add_runs_with_bold(p, text)


def add_page_footer():
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AI 求职智能匹配智能体 · 方案说明")
    r.font.size = Pt(8)
    r.font.color.rgb = GRAY
    set_cn_font(r)


lines = open(SRC, encoding="utf-8").read().split("\n")
order_idx = 0
for line in lines:
    line = line.rstrip()
    if not line:
        continue
    if line.startswith("# "):
        add_cover(line[2:], "问题诊断 · 方案设计 · AI 选型 · 关键配置 · 迭代记录 · 效果评估")
    elif line.startswith("## "):
        order_idx = 0
        add_h1(line[3:])
    elif re.match(r"^\d+\.\s", line):
        order_idx += 1
        content = re.sub(r"^\d+\.\s", "", line)
        add_list_item(content, ordered=True, idx=order_idx)
    elif line.startswith("- "):
        add_list_item(line[2:], ordered=False)
    else:
        add_body(line)

add_page_footer()
doc.save(OUT)
print("saved", OUT)
